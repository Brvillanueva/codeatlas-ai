"""Editable Word report generation from deterministic repository analysis."""

from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from codeatlas.exceptions import OutputAlreadyExistsError
from codeatlas.models import ArchitectureNarrative, FileAnalysis, FunctionInfo, RepositoryAnalysis

_BLUE = RGBColor(46, 116, 181)
_DARK_BLUE = RGBColor(31, 77, 120)
_MUTED = RGBColor(89, 89, 89)
_HEADER_FILL = "F2F4F7"
_DETAILED_MODULE_LIMIT = 25


class WordReportGenerator:
    """Create a concise, editable architecture report in DOCX format."""

    def write(
        self,
        analysis: RepositoryAnalysis,
        output_path: Path,
        force: bool = False,
        narrative: ArchitectureNarrative | None = None,
    ) -> Path:
        output = output_path.expanduser().resolve()
        if output.exists() and not force:
            raise OutputAlreadyExistsError(
                f"Output already exists: {output}. Use --force to overwrite it."
            )

        output.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        self._configure_document(document)
        self._add_cover(document, analysis)
        document.add_page_break()
        self._add_contents(document, includes_ai=narrative is not None)
        document.add_section(WD_SECTION_START.NEW_PAGE)
        self._add_footer(document.sections[-1])
        self._add_report(document, analysis, narrative)
        document.save(output)
        return output

    def _configure_document(self, document: Document) -> None:
        section = document.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

        normal = document.styles["Normal"]
        normal.font.name = "Calibri"
        normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.1

        for name, size, color, before, after in (
            ("Heading 1", 16, _BLUE, 16, 8),
            ("Heading 2", 13, _BLUE, 12, 6),
            ("Heading 3", 12, _DARK_BLUE, 8, 4),
        ):
            style = document.styles[name]
            style.font.name = "Calibri"
            style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
            style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
            style.font.size = Pt(size)
            style.font.color.rgb = color
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)

        code_style = document.styles.add_style("CodeAtlas Code", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Consolas"
        code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        code_style.font.size = Pt(8.5)
        code_style.paragraph_format.space_after = Pt(4)

        for item in document.sections:
            self._add_footer(item)

    def _add_cover(self, document: Document, analysis: RepositoryAnalysis) -> None:
        document.add_paragraph()
        label = document.add_paragraph()
        label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = label.add_run("CODEATLAS AI")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = _BLUE

        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(75)
        title.paragraph_format.space_after = Pt(10)
        run = title.add_run("Informe de análisis técnico")
        run.bold = True
        run.font.size = Pt(30)
        run.font.color.rgb = _DARK_BLUE

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(36)
        run = subtitle.add_run("Evidencia estática de un repositorio Python")
        run.font.size = Pt(15)
        run.font.color.rgb = _MUTED

        metadata = self._table(document, 3, 2, (Inches(1.6), Inches(4.9)))
        for row, label_text, value in zip(
            metadata.rows,
            ("Proyecto", "Ruta analizada", "Generado"),
            (
                analysis.repository_name,
                analysis.repository_path,
                datetime.now().strftime("%Y-%m-%d %H:%M"),
            ),
            strict=True,
        ):
            row.cells[0].text = label_text
            row.cells[1].text = value
            self._shade_cell(row.cells[0], _HEADER_FILL)
            for paragraph in row.cells[0].paragraphs:
                paragraph.runs[0].bold = True

        note = document.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note.paragraph_format.space_before = Pt(45)
        note_run = note.add_run(
            "Este informe se genera a partir de la estructura y dependencias detectadas "
            "en el código. No sustituye una revisión humana del dominio de negocio."
        )
        note_run.italic = True
        note_run.font.size = Pt(9.5)
        note_run.font.color.rgb = _MUTED

    def _add_contents(self, document: Document, includes_ai: bool) -> None:
        document.add_heading("Contenido", level=1)
        items = [
            "1. Como leer este informe",
            "2. Resumen para orientar la lectura",
            "3. Hipótesis de organización técnica",
            "4. Datos detectados y alcance",
            "5. Tecnologías y dependencias externas",
            "6. Áreas para revisión",
            "Anexo técnico",
        ]
        if includes_ai:
            items.insert(6, "7. Interpretación generada por IA (opcional)")
        for item in items:
            document.add_paragraph(item)

    def _add_report(
        self,
        document: Document,
        analysis: RepositoryAnalysis,
        narrative: ArchitectureNarrative | None,
    ) -> None:
        source_files = self._source_files(analysis.files)
        test_files = [file for file in analysis.files if self._is_test_file(file.path)]
        initializer_files = [
            file
            for file in analysis.files
            if self._is_initializer(file.path) and not self._is_test_file(file.path)
        ]
        source_paths = {file.path for file in source_files}
        central_modules = [
            module for module in analysis.central_modules if module in source_paths
        ]

        document.add_heading("1. Como leer este informe", level=1)
        document.add_paragraph(
            "Este informe separa los hechos detectados de las hipótesis y, cuando se solicite, "
            "de la interpretación generada por IA. Ninguna de las tres capas sustituye una "
            "revisión "
            "humana del dominio de negocio."
        )
        self._add_reading_legend(document)

        document.add_heading("2. Resumen para orientar la lectura", level=1)
        stats = analysis.stats
        document.add_paragraph(
            "CodeAtlas AI analizó "
            f"{stats.python_files} archivo(s) Python dentro de {stats.files_discovered} archivo(s) "
            f"descubierto(s). Esta vista principal se concentra en {len(source_files)} módulo(s) "
            f"de código y resume {len(test_files)} archivo(s) de prueba. El resultado sirve como "
            "punto de partida para comprender el código, "
            "planear una revisión técnica y detectar zonas que merecen atención, pero no confirma "
            "por sí solo cómo funciona el sistema en ejecución."
        )
        summary = self._table(document, 3, 3, (Inches(2.15), Inches(2.15), Inches(2.2)))
        for cell, label, value in zip(
            summary.rows[0].cells,
            ("Estructura", "Código", "Dependencias"),
            (
                f"{len(source_files)} módulo(s) principal(es)",
                f"{sum(len(file.classes) for file in source_files)} clase(s) y "
                f"{sum(len(file.functions) for file in source_files)} función(es)",
                f"{stats.internal_imports} interna(s) detectada(s)",
            ),
            strict=True,
        ):
            cell.text = label
            self._shade_cell(cell, _HEADER_FILL)
            cell.paragraphs[0].runs[0].bold = True
        for cell, value in zip(
            summary.rows[1].cells,
            ("Métodos", "Módulos centrales", "Errores de análisis"),
            strict=True,
        ):
            cell.text = value
            self._shade_cell(cell, "FAFAFA")
            cell.paragraphs[0].runs[0].bold = True
        for cell, value in zip(
            summary.rows[2].cells,
            (
                str(stats.methods),
                str(len(central_modules)),
                str(stats.files_with_errors),
            ),
            strict=True,
        ):
            cell.text = value

        document.add_heading("3. Hipótesis de organización técnica", level=1)
        self._add_architecture_overview(document, analysis, source_paths)

        document.add_heading("4. Datos detectados y alcance", level=1)
        document.add_paragraph(
            "El informe utiliza análisis estático: nombres de archivos y módulos, estructuras "
            "Python, importaciones y relaciones internas. No ejecuta el programa, no inspecciona "
            "bases de datos "
            "ni confirma reglas de negocio que no estén expresadas en el código."
        )
        self._add_label_detail_table(
            document,
            (
                ("Repositorio", analysis.repository_name),
                ("Ruta", analysis.repository_path),
                ("Archivos Python", str(stats.python_files)),
                ("Límites", "Solo los archivos que el escáner pudo leer y analizar."),
            ),
        )

        document.add_heading("Módulos que requieren atención técnica", level=2)
        if central_modules:
            document.add_paragraph(
                "Los siguientes módulos son importados por más archivos dentro del conjunto "
                "analizado. "
                "Conviene revisarlos primero al estudiar cambios técnicos, porque otros módulos "
                "dependen de ellos:")
            for module in central_modules:
                document.add_paragraph(module, style="List Bullet")
        else:
            document.add_paragraph(
                "No se detectaron relaciones internas suficientes para identificar módulos "
                "centrales."
            )

        source_dependencies = [
            dependency
            for dependency in analysis.internal_dependencies
            if dependency.source in source_paths and dependency.target in source_paths
        ]
        if source_dependencies:
            table = self._table(
                document,
                len(source_dependencies) + 1,
                3,
                (Inches(2.15), Inches(2.15), Inches(2.2)),
            )
            self._set_header(table.rows[0].cells, ("Origen", "Destino", "Importación"))
            for row, dependency in zip(table.rows[1:], source_dependencies, strict=True):
                row.cells[0].text = dependency.source
                row.cells[1].text = dependency.target
                row.cells[2].text = dependency.imported_module
        else:
            document.add_paragraph(
                "No se detectaron dependencias internas entre los módulos analizados."
            )

        document.add_heading("5. Tecnologías y dependencias externas", level=1)
        self._add_import_summary(document, source_files)

        document.add_heading("6. Áreas para revisión", level=1)
        recommendations = [
            "Validar con el equipo la responsabilidad real de los módulos más importados "
            "antes de realizar cambios.",
            "Revisar las dependencias externas frente al archivo de paquetes del proyecto.",
            "Usar este informe junto con pruebas y revisión humana; el análisis estático no "
            "ejecuta la aplicación.",
        ]
        if analysis.errors:
            recommendations.insert(
                0,
                "Corregir primero los archivos con errores de lectura o sintaxis para obtener "
                "una cobertura mayor.",
            )
        for recommendation in recommendations:
            document.add_paragraph(recommendation, style="List Bullet")

        if narrative:
            self._add_ai_narrative(document, narrative)

        document.add_heading("Anexo técnico", level=1)
        document.add_heading("A. Inventario de código principal", level=2)
        inventory = self._table(
            document,
            len(source_files) + 1,
            5,
            (Inches(1.6), Inches(1.25), Inches(1.2), Inches(1.2), Inches(1.25)),
        )
        self._set_header(
            inventory.rows[0].cells,
            ("Archivo", "Módulo", "Clases", "Funciones", "Estado"),
        )
        for row, file in zip(inventory.rows[1:], source_files, strict=True):
            row.cells[0].text = file.path
            row.cells[1].text = file.module_name
            row.cells[2].text = str(len(file.classes))
            row.cells[3].text = str(len(file.functions))
            row.cells[4].text = "Error" if file.syntax_error else "Analizado"

        document.add_heading("B. Detalle de módulos seleccionados", level=2)
        detailed_files = self._detailed_files(source_files, central_modules)
        if len(detailed_files) < len(source_files):
            document.add_paragraph(
                f"Para conservar el informe legible, este anexo muestra {len(detailed_files)} de "
                f"{len(source_files)} módulos principales. El comando 'codeatlas analyze' permite "
                "exportar el inventario completo en JSON."
            )
        for file in detailed_files:
            self._add_file_detail(document, file)

        document.add_heading("C. Pruebas e inicializadores resumidos", level=2)
        document.add_paragraph(
            f"Se excluyeron {len(test_files)} archivo(s) de prueba y {len(initializer_files)} "
            "inicializador(es) de la vista principal para evitar que oculten la estructura del "
            "código."
        )
        if test_files:
            document.add_paragraph(
                "Pruebas detectadas: " + ", ".join(file.path for file in test_files[:10])
            )

        document.add_heading("D. Diagramas técnicos disponibles", level=2)
        document.add_paragraph(
            "El mapa de dependencias y el diagrama de clases se exportan por separado. "
            "No se incluye un diagrama de arquitectura automático si la evidencia estática no "
            "es suficiente."
        )

        document.add_heading("E. Errores y limitaciones", level=2)
        if analysis.errors:
            errors = self._table(
                document,
                len(analysis.errors) + 1,
                3,
                (Inches(1.75), Inches(1.3), Inches(3.45)),
            )
            self._set_header(errors.rows[0].cells, ("Archivo", "Tipo", "Detalle"))
            for row, error in zip(errors.rows[1:], analysis.errors, strict=True):
                row.cells[0].text = error.path
                row.cells[1].text = error.kind
                row.cells[2].text = error.message
        else:
            document.add_paragraph(
                "No se registraron errores de lectura ni sintaxis durante el análisis."
            )
        document.add_paragraph(
            "Limitaciones: el resultado describe el código de forma estática. Las llamadas "
            "dinámicas, configuraciones externas, archivos no Python, comportamiento en "
            "ejecución y reglas de negocio "
            "requieren validación adicional."
        )

    def _add_architecture_overview(
        self, document: Document, analysis: RepositoryAnalysis, source_paths: set[str]
    ) -> None:
        document.add_paragraph(
            "Esta clasificación es heurística: usa la ruta y los nombres de archivos. "
            "No representa por sí sola una arquitectura validada."
        )
        components = [
            (item, [path for path in item.files if path in source_paths])
            for item in analysis.architecture.components
            if item.role != "tests"
        ]
        components = [item for item in components if item[1]]
        if components:
            table = self._table(
                document,
                len(components) + 1,
                4,
                (Inches(1.5), Inches(1.0), Inches(1.0), Inches(3.0)),
            )
            self._set_header(
                table.rows[0].cells,
                ("Componente", "Archivos", "Confianza", "Evidencia"),
            )
            for row, (component, files) in zip(table.rows[1:], components, strict=True):
                row.cells[0].text = component.name
                row.cells[1].text = str(len(files))
                row.cells[2].text = f"{component.confidence:.0%}"
                row.cells[3].text = "; ".join(component.evidence[:2])
        relationships = [
            dependency
            for dependency in analysis.architecture.dependencies
            if any(path in source_paths for path in dependency.evidence_files)
        ]
        if relationships:
            document.add_paragraph("Relaciones entre componentes detectadas:")
            for dependency in relationships:
                document.add_paragraph(
                    f"{dependency.source} -> {dependency.target} "
                    f"({dependency.imports_count} import(s); evidencia: "
                    f"{', '.join(dependency.evidence_files)})",
                    style="List Bullet",
                )
        else:
            document.add_paragraph("No se detectaron relaciones entre componentes diferentes.")

    def _add_import_summary(self, document: Document, files: list[FileAnalysis]) -> None:
        categories: defaultdict[str, set[str]] = defaultdict(set)
        for file in files:
            for item in file.imports:
                if item.classification is None:
                    continue
                root = item.module.split(".", maxsplit=1)[0] if item.module else ""
                if root:
                    categories[item.classification].add(root)
        third_party = sorted(categories["third_party"])
        standard_library = sorted(categories["standard_library"])
        unresolved_internal = sorted(categories["unresolved_internal"])
        self._add_label_detail_table(
            document,
            (
                (
                    "Terceros",
                    ", ".join(third_party)
                    if third_party
                    else "No se detectaron paquetes de terceros en el código principal.",
                ),
                (
                    "Biblioteca estándar",
                    f"{len(standard_library)} módulo(s) detectado(s). No requieren instalación "
                    "externa.",
                ),
                (
                    "Internos sin resolver",
                    ", ".join(unresolved_internal)
                    if unresolved_internal
                    else "No se detectaron imports internos sin resolver.",
                ),
            ),
        )
        if unresolved_internal:
            document.add_paragraph(
                "Los imports internos sin resolver pueden indicar una ruta incompleta o un "
                "archivo fuera del alcance analizado. Deben revisarse antes de tratarlos "
                "como paquetes externos."
            )

    def _source_files(self, files: list[FileAnalysis]) -> list[FileAnalysis]:
        return [
            file
            for file in files
            if not self._is_test_file(file.path) and not self._is_initializer(file.path)
        ]

    def _detailed_files(
        self, files: list[FileAnalysis], central_modules: list[str]
    ) -> list[FileAnalysis]:
        central = set(central_modules)
        return sorted(
            files,
            key=lambda file: (
                file.path not in central,
                -(file.line_count + len(file.imports) * 5 + len(file.classes) * 10),
                file.path,
            ),
        )[:_DETAILED_MODULE_LIMIT]

    def _is_test_file(self, path: str) -> bool:
        parts = path.replace("\\", "/").split("/")
        return "tests" in parts or parts[-1].startswith("test_")

    def _is_initializer(self, path: str) -> bool:
        return path.replace("\\", "/").endswith("/__init__.py") or path == "__init__.py"

    def _truncate(self, text: str, limit: int = 280) -> str:
        normalized = " ".join(text.strip().split())
        return normalized if len(normalized) <= limit else f"{normalized[: limit - 1]}…"

    def _add_reading_legend(self, document: Document) -> None:
        table = self._table(
            document,
            4,
            3,
            (Inches(1.5), Inches(2.45), Inches(2.65)),
        )
        self._set_header(table.rows[0].cells, ("Etiqueta", "Qué significa", "Cómo usarla"))
        rows = (
            (
                "Dato detectado",
                "Archivo, import, clase, función o error extraído localmente.",
                "Se puede verificar en el código analizado.",
            ),
            (
                "Hipótesis estática",
                "Clasificación basada en rutas, nombres y relaciones de importación.",
                "Debe validarse con quien conoce el proyecto.",
            ),
            (
                "Interpretación IA",
                "Explicación opcional generada a partir de evidencia seleccionada.",
                "Nunca se presenta como hecho; incluye archivos de respaldo.",
            ),
        )
        fills = ("EAF4E3", "FFF3CD", "E8EEF5")
        for row, values, fill in zip(table.rows[1:], rows, fills, strict=True):
            for cell, value in zip(row.cells, values, strict=True):
                cell.text = value
            self._shade_cell(row.cells[0], fill)
            row.cells[0].paragraphs[0].runs[0].bold = True

    def _add_ai_narrative(
        self, document: Document, narrative: ArchitectureNarrative
    ) -> None:
        document.add_heading("7. Interpretación generada por IA (opcional)", level=1)
        document.add_paragraph(
            "Esta seccion es una interpretacion, no un dato comprobado por si sola. "
            "Debe leerse junto con la evidencia indicada."
        )
        document.add_heading("Proposito probable", level=3)
        document.add_paragraph(narrative.purpose)
        if narrative.component_responsibilities:
            document.add_heading("Responsabilidades sugeridas", level=3)
            self._add_label_detail_table(
                document, tuple(sorted(narrative.component_responsibilities.items()))
            )
        for heading, values in (
            ("Flujo principal", narrative.main_flow),
            ("Riesgos iniciales", narrative.risks),
            ("Preguntas abiertas", narrative.open_questions),
        ):
            if values:
                document.add_heading(heading, level=3)
                for value in values:
                    document.add_paragraph(value, style="List Bullet")
        if narrative.evidence:
            document.add_heading("Evidencia citada por la IA", level=3)
            for item in narrative.evidence:
                document.add_paragraph(
                    f"{item.claim} - archivos: {', '.join(item.files)}", style="List Bullet"
                )

    def _add_file_detail(self, document: Document, file: FileAnalysis) -> None:
        document.add_heading(f"Archivo: {file.path}", level=3)
        facts = [
            f"Módulo: {file.module_name}",
            f"Líneas detectadas: {file.line_count}",
            f"Clases: {len(file.classes)}; funciones de módulo: {len(file.functions)}; "
            f"imports: {len(file.imports)}.",
        ]
        if file.docstring:
            facts.append(f"Descripción declarada: {self._truncate(file.docstring)}")
        if file.syntax_error:
            facts.append(f"Error de sintaxis: {file.syntax_error}")
        for fact in facts:
            document.add_paragraph(fact)

        if file.classes:
            document.add_paragraph("Clases detectadas:")
            for item in file.classes:
                bases = f" ({', '.join(item.bases)})" if item.bases else ""
                document.add_paragraph(f"{item.name}{bases}", style="List Bullet")
                for method in item.methods:
                    document.add_paragraph(
                        self._signature(method), style="List Bullet 2"
                    )
        if file.functions:
            document.add_paragraph("Funciones de módulo detectadas:")
            for function in file.functions:
                document.add_paragraph(self._signature(function), style="List Bullet")
        if file.imports:
            document.add_paragraph("Imports detectados:")
            statuses = {
                "internal": "interna",
                "standard_library": "biblioteca estándar",
                "third_party": "tercero",
                "unresolved_internal": "interna sin resolver",
            }
            for item in file.imports:
                names = ", ".join(item.imported_names)
                suffix = f" ({names})" if names else ""
                status = statuses.get(item.classification, "sin resolver")
                document.add_paragraph(f"{item.module}{suffix} - {status}", style="List Bullet")

    def _signature(self, function: FunctionInfo) -> str:
        parameters = []
        for parameter in function.parameters:
            value = parameter.name
            if parameter.annotation:
                value += f": {parameter.annotation}"
            if parameter.default:
                value += f" = {parameter.default}"
            parameters.append(value)
        prefix = "async " if function.is_async else ""
        result = f"{prefix}{function.name}({', '.join(parameters)})"
        if function.return_type:
            result += f" -> {function.return_type}"
        return result

    def _add_label_detail_table(
        self, document: Document, rows: tuple[tuple[str, str], ...]
    ) -> None:
        table = self._table(document, len(rows), 2, (Inches(1.85), Inches(4.65)))
        for row, (label, value) in zip(table.rows, rows, strict=True):
            row.cells[0].text = label
            row.cells[1].text = value
            self._shade_cell(row.cells[0], _HEADER_FILL)
            row.cells[0].paragraphs[0].runs[0].bold = True

    def _table(self, document: Document, rows: int, columns: int, widths: tuple[Inches, ...]):
        table = document.add_table(rows=rows, cols=columns)
        table.style = "Table Grid"
        table.autofit = False
        widths_dxa = tuple(int(width.twips) for width in widths)
        properties = table._tbl.tblPr
        table_width = properties.find(qn("w:tblW"))
        table_width.set(qn("w:w"), str(sum(widths_dxa)))
        table_width.set(qn("w:type"), "dxa")
        table_indent = properties.find(qn("w:tblInd"))
        if table_indent is None:
            table_indent = OxmlElement("w:tblInd")
            properties.append(table_indent)
        table_indent.set(qn("w:w"), "120")
        table_indent.set(qn("w:type"), "dxa")
        for grid_column, width in zip(table._tbl.tblGrid.gridCol_lst, widths_dxa, strict=True):
            grid_column.set(qn("w:w"), str(width))
        for row in table.rows:
            for cell, width, width_dxa in zip(row.cells, widths, widths_dxa, strict=True):
                cell.width = width
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                cell_width.set(qn("w:w"), str(width_dxa))
                cell_width.set(qn("w:type"), "dxa")
                self._set_cell_margins(cell)
        return table

    def _set_header(self, cells: tuple | list, values: tuple[str, ...]) -> None:
        for cell, value in zip(cells, values, strict=True):
            cell.text = value
            self._shade_cell(cell, _HEADER_FILL)
            cell.paragraphs[0].runs[0].bold = True

    def _shade_cell(self, cell, fill: str) -> None:
        properties = cell._tc.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        properties.append(shading)

    def _set_cell_margins(self, cell) -> None:
        properties = cell._tc.get_or_add_tcPr()
        margins = properties.first_child_found_in("w:tcMar")
        if margins is None:
            margins = OxmlElement("w:tcMar")
            properties.append(margins)
        for side in ("top", "start", "bottom", "end"):
            node = margins.find(qn(f"w:{side}"))
            if node is None:
                node = OxmlElement(f"w:{side}")
                margins.append(node)
            node.set(qn("w:w"), "80" if side in {"top", "bottom"} else "120")
            node.set(qn("w:type"), "dxa")

    def _add_footer(self, section) -> None:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run("CodeAtlas AI | Informe generado automáticamente")
        run.font.size = Pt(8)
        run.font.color.rgb = _MUTED

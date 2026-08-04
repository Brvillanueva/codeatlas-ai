from pathlib import Path

from docx import Document
from typer.testing import CliRunner

from codeatlas.cli import app

runner = CliRunner()
FIXTURE_ROOT = Path(__file__).parent / "fixtures" / "simple_project"


def test_help_is_available() -> None:
    result = runner.invoke(app, ["--help"])

    assert result.exit_code == 0
    assert "inspect" in result.stdout
    assert "analyze" in result.stdout
    assert "graph" in result.stdout
    assert "dependencies" in result.stdout
    assert "classes" in result.stdout
    assert "report" in result.stdout


def test_inspect_displays_summary() -> None:
    result = runner.invoke(app, ["inspect", str(FIXTURE_ROOT)])

    assert result.exit_code == 0
    assert "CodeAtlas AI" in result.stdout
    assert "completed" in result.stdout
    assert "Python files" in result.stdout


def test_analyze_exports_json(tmp_path: Path) -> None:
    output = tmp_path / "analysis.json"
    result = runner.invoke(app, ["analyze", str(FIXTURE_ROOT), "--output", str(output)])

    assert result.exit_code == 0
    assert output.exists()
    assert '"repository_name"' in output.read_text(encoding="utf-8")


def test_graph_defaults_to_technical_view(tmp_path: Path) -> None:
    output = tmp_path / "graph.mmd"
    result = runner.invoke(app, ["graph", str(FIXTURE_ROOT), "--output", str(output)])

    assert result.exit_code == 0
    assert output.read_text(encoding="utf-8").startswith("graph TD")


def test_graph_can_export_complete_technical_view(tmp_path: Path) -> None:
    output = tmp_path / "graph.mmd"
    result = runner.invoke(
        app, ["graph", str(FIXTURE_ROOT), "--output", str(output), "--view", "technical"]
    )

    assert result.exit_code == 0
    assert output.read_text(encoding="utf-8").startswith("graph TD")


def test_dependencies_exports_a_package_summary_by_default(tmp_path: Path) -> None:
    output = tmp_path / "dependencies.mmd"
    result = runner.invoke(app, ["dependencies", str(FIXTURE_ROOT), "--output", str(output)])

    assert result.exit_code == 0
    assert output.read_text(encoding="utf-8").startswith("flowchart LR")


def test_dependencies_can_export_complete_file_imports(tmp_path: Path) -> None:
    output = tmp_path / "dependencies.mmd"
    result = runner.invoke(
        app, ["dependencies", str(FIXTURE_ROOT), "--level", "file", "--output", str(output)]
    )

    assert result.exit_code == 0
    assert output.read_text(encoding="utf-8").startswith("flowchart LR")


def test_dependencies_accepts_a_package_scope(tmp_path: Path) -> None:
    output = tmp_path / "dependencies.mmd"
    result = runner.invoke(
        app,
        ["dependencies", str(FIXTURE_ROOT), "--package", "package", "--output", str(output)],
    )

    assert result.exit_code == 0
    assert "package" in output.read_text(encoding="utf-8")


def test_classes_exports_a_class_diagram(tmp_path: Path) -> None:
    output = tmp_path / "classes.mmd"
    result = runner.invoke(app, ["classes", str(FIXTURE_ROOT), "--output", str(output)])

    assert result.exit_code == 0
    assert output.read_text(encoding="utf-8").startswith("classDiagram")


def test_report_exports_editable_word_document(tmp_path: Path) -> None:
    output = tmp_path / "report.docx"
    result = runner.invoke(app, ["report", str(FIXTURE_ROOT), "--output", str(output)])

    assert result.exit_code == 0
    assert output.exists()
    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    assert "Contenido" in text
    assert "Hipótesis de organización técnica" in text
    assert "Datos detectados y alcance" in text
    assert "Archivo: tests/test_service.py" not in text
    assert "Dato detectado" in table_text
    assert "Hipótesis estática" in table_text


def test_report_ai_without_key_fails_before_creating_document(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    output = tmp_path / "report.docx"

    result = runner.invoke(app, ["report", str(FIXTURE_ROOT), "--ai", "--output", str(output)])

    assert result.exit_code == 1
    assert not output.exists()
